from __future__ import annotations

import argparse
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Emu
from docx.text.paragraph import Paragraph


def main() -> None:
    parser = argparse.ArgumentParser(description="用高分辨率 PNG 替换周报中的原生图1")
    parser.add_argument("docx", type=Path)
    parser.add_argument("png", type=Path)
    args = parser.parse_args()

    docx_path = args.docx.resolve()
    png_path = args.png.resolve()
    document = Document(docx_path)

    matches = []
    for inline in document.element.xpath(".//wp:inline"):
        if inline.xpath('.//*[local-name()="chart"]'):
            matches.append(inline)
    if len(matches) != 1:
        raise RuntimeError(f"预期找到1个原生图表，实际找到{len(matches)}个")

    inline = matches[0]
    extent = inline.xpath("./wp:extent")[0]
    width = Emu(int(extent.get("cx")))
    height = Emu(int(extent.get("cy")))
    drawing = inline.getparent()
    old_run = drawing.getparent()
    paragraph_element = old_run.getparent()
    paragraph = Paragraph(paragraph_element, document)
    paragraph.add_run().add_picture(str(png_path), width=width, height=height)
    paragraph_element.remove(old_run)

    with tempfile.NamedTemporaryFile(
        dir=docx_path.parent, prefix="figure1_image_", suffix=".docx", delete=False
    ) as handle:
        temp_path = Path(handle.name)
    try:
        document.save(temp_path)
        shutil.move(str(temp_path), str(docx_path))
    finally:
        if temp_path.exists():
            temp_path.unlink()
    print(f"已替换图1，显示尺寸 {width.inches:.3f} x {height.inches:.3f} 英寸")


if __name__ == "__main__":
    main()
